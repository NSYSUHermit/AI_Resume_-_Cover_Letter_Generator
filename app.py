import streamlit as st
import google.generativeai as genai
import jinja2
import subprocess
import os
import json
import tempfile
import shutil
import base64
import io
import difflib
from docx import Document
from docx.shared import Pt, Inches
import streamlit.components.v1 as components
import streamlit_ace as st_ace # 引入 streamlit-ace 套件
from datetime import datetime
from firebase_dashboard import init_firebase, authenticate_user, register_user, render_dashboard, save_application, render_interview_progress, save_user_profile, load_user_profile, predict_interview_questions, analyze_skill_gap

# ---------------------------------------------------------
# 初始化 Session State (JSON 資料結構)
# ---------------------------------------------------------
if "resume_data" not in st.session_state:
    st.session_state.resume_data = {
        "heading": {
            "name": "John Doe",
            "email": "johndoe@example.com",
            "phone": "+1-234-567-8900",
            "website": "github.com/johndoe",
            "linkedin": "linkedin.com/in/johndoe"
        },
        "cover_letter":"I am writing to express my strong interest in the AI Engineer position at Google. I can rapidly master the intricacies of semiconductor equipment, contributing significantly to your team’s success.",
        "target_company":"Google",
        "target_role":"AI Engineer",
        "about me more": "I have independently handled end-to-end development, from low-level model optimization and backend microservice architecture to frontend user interfaces...",
        "summary": "Software Engineer with 3 years of experience specializing in scalable backend architectures and AI-driven systems...",
        "education": [
            {
                "degree": "Master of Science in Computer Science",
                "time_period": "Aug 2021 - May 2023",
                "school": "State University",
                "school_location": "New York, NY"
            }
        ],
        "experience": [
            {
                "role": "Software Engineer",
                "team": "Backend Core Team",
                "company": "Tech Corp",
                "company_location": "San Francisco, CA",
                "time_duration": "Jun 2023 - Present",
                "details": [
                    {
                        "title": "Microservices Architecture",
                        "description": "Architected a high-performance backend using Python and FastAPI to orchestrate automated workflows, improving efficiency by 40%."
                    }
                ]
            }
        ],
        "projects": [
            {
                "name": "E-Commerce Platform",
                "time": "Jan 2023 - May 2023",
                "description": "Led a team of 4 to develop a full-stack e-commerce web application using React and Django."
            }
        ],
        "patents": [],
        "skills": {
            "set1": {
                "title": "Backend & Architecture",
                "items": ["Python (FastAPI, Django)", "RESTful API Design", "Docker", "AWS"]
            }
        }
    }

if "ai_report" not in st.session_state:
    st.session_state.ai_report = ""

if "optimized_resume_data" not in st.session_state:
    st.session_state.optimized_resume_data = None

if "base_resume_saved_time" not in st.session_state:
    st.session_state.base_resume_saved_time = None
if "optimized_resume_saved_time" not in st.session_state:
    st.session_state.optimized_resume_saved_time = None
if "base_editor_key" not in st.session_state:
    st.session_state.base_editor_key = 0
if "opt_editor_key" not in st.session_state:
    st.session_state.opt_editor_key = 0
if "ats_metrics" not in st.session_state:
    st.session_state.ats_metrics = None
if "changelog" not in st.session_state:
    st.session_state.changelog = ""
if "custom_prompt" not in st.session_state:
    st.session_state.custom_prompt = """You are an elite Career Strategist and ATS Architect. Overhaul the resume based on the JD using these rules:

1. **Strategic Quantization**: Every experience bullet point MUST include a metric (%, $, time saved, scale). Transform tasks into results. (e.g., instead of 'Used microservices', use 'Architected 20+ FastAPI microservices, reducing latency by 45%').
2. **Aggressive Action Verbs**: Use high-ownership verbs like 'Spearheaded', 'Engineered', 'Orchestrated', and 'Optimized'.
3. **ATS Semantic Mapping**: Perform 'Horizontal Shifts'. If the JD demands GCP but I have AWS, frame it as 'Cloud-native architecture expertise (AWS/GCP)'. Ensure high keyword density in Summary and Skills.
4. **Pain Point Alignment**: Identify the 3 biggest problems mentioned in the JD and ensure the Summary and top Experience points explicitly demonstrate how I solved those exact problems before.
5. **Authoritative Tone**: Maintain a bold, professional, and result-oriented voice that positions me as the 'perfect solution' to the company's needs."""
if "api_key" not in st.session_state:
    st.session_state.api_key = ""
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user_email" not in st.session_state:
    st.session_state.user_email = ""
if "preview_pdf_bytes" not in st.session_state:
    st.session_state.preview_pdf_bytes = None
if "resume_preview_bytes" not in st.session_state:
    st.session_state.resume_preview_bytes = None
if "cover_letter_preview_bytes" not in st.session_state:
    st.session_state.cover_letter_preview_bytes = None
if "cover_letter_filename" not in st.session_state:
    st.session_state.cover_letter_filename = "cover_letter.pdf"
if 'resume_dl_data' not in st.session_state:
    st.session_state.resume_dl_data = None
if 'cl_dl_data' not in st.session_state:
    st.session_state.cl_dl_data = None

# ---------------------------------------------------------
# AI 核心邏輯 (ATS 關鍵字分析與履歷優化)
# ---------------------------------------------------------
def parse_pdf_resume_to_json(pdf_bytes, api_key):
    if not api_key:
        return False, "⚠️ Error: Please set your GEMINI API KEY in the sidebar first.", None

    try:
        genai.configure(api_key=api_key)
        model_name = st.session_state.get("ai_model", "gemini-1.5-flash")
        model = genai.GenerativeModel(model_name)

        prompt = """
        You are an expert resume parser. I will provide you with a resume in PDF format.
        Extract all the relevant information and structure it STRICTLY in the following JSON format.
        Return ONLY valid JSON.
        """

        pdf_part = {"mime_type": "application/pdf", "data": pdf_bytes}
        response = model.generate_content([prompt, pdf_part])
        raw_text = response.text.strip()
        if "```" in raw_text:
            raw_text = raw_text.split("```")[1]
            if raw_text.startswith("json"): raw_text = raw_text[4:]
        
        parsed_json = json.loads(raw_text.strip())
        return True, "✅ Success!", parsed_json
    except Exception as e:
        return False, f"⚠️ Error during PDF parsing: {e}", None

def build_optimization_prompt(jd_text, custom_prompt, enable_ats, check_visa, resume_data):
    ats_example = ""
    if enable_ats:
        ats_example = '"keyword_analysis": {"jd_keywords": [], "original_hits": [], "optimized_hits": [], "newly_added": [], "missing_keywords": []},'

    return f"""
You are a Resume Expert. Optimize the [Original Resume] based on the [Target JD].
[COMMANDS]: {custom_prompt}
[RULES]: 
1. Return ONLY valid JSON.
2. { "- Step 1: Check visa restrictions." if check_visa else "" }
3. Generate 'cover_letter' and 'optimized_resume'.

[Target JD]: {jd_text}
[Original Resume JSON]: {json.dumps(resume_data, ensure_ascii=False)}

[FORMAT]:
{{
    "visa_blocked": false, "reason": "", "changelog": "",
    {ats_example}
    "optimized_resume": {{ ...STRUCTURE... }}
}}
"""

def ai_optimize_and_update(jd_text, custom_prompt, enable_ats, check_visa):
    try:
        api_key = st.session_state.get("api_key", "")
        if not api_key: return False, "Missing API Key."
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(st.session_state.get("ai_model", "gemini-1.5-flash"))
        
        final_prompt = build_optimization_prompt(jd_text, custom_prompt, enable_ats, check_visa, st.session_state.resume_data)
        response = model.generate_content(final_prompt)
        
        raw_text = response.text.strip()
        if "```" in raw_text:
            raw_text = raw_text.split("```")[1]
            if raw_text.startswith("json"): raw_text = raw_text[4:]
        
        ai_result = json.loads(raw_text.strip())
        
        if ai_result.get("visa_blocked"):
            return False, f"Visa Restriction: {ai_result.get('reason')}"
            
        st.session_state.optimized_resume_data = ai_result.get("optimized_resume")
        st.session_state.opt_editor_key += 1
        st.session_state.changelog = ai_result.get("changelog", "")
        
        if enable_ats and "keyword_analysis" in ai_result:
            kw = ai_result["keyword_analysis"]
            tot = max(1, len(kw.get("optimized_hits", [])) + len(kw.get("missing_keywords", [])))
            st.session_state.ats_metrics = {
                "total": tot,
                "original_count": len(kw.get("original_hits", [])),
                "optimized_count": len(kw.get("optimized_hits", [])),
                "original_pct": int((len(kw.get("original_hits", []))/tot)*100),
                "optimized_pct": int((len(kw.get("optimized_hits", []))/tot)*100),
                "optimized_hits": kw.get("optimized_hits", []),
                "newly_added": kw.get("newly_added", []),
                "missing_keywords": kw.get("missing_keywords", [])
            }
        return True, "✅ Success"
    except Exception as e:
        return False, str(e)

def predict_interview_questions(jd_text, resume_data):
    try:
        api_key = st.session_state.get("api_key", "")
        if not api_key: return None
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(st.session_state.get("ai_model", "gemini-1.5-flash"))
        prompt = f"Predict interview questions for this JD and Resume. Return JSON with 'technical' and 'behavioral' lists. [JD]: {jd_text} [Resume]: {json.dumps(resume_data)}"
        response = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(response_mime_type="application/json"))
        return json.loads(response.text)
    except: return None

def analyze_skill_gap(jd_text, resume_data):
    try:
        api_key = st.session_state.get("api_key", "")
        if not api_key: return None
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(st.session_state.get("ai_model", "gemini-1.5-flash"))
        prompt = f"Analyze skill gap (radar chart data). Return JSON with 'categories', 'candidate_scores', 'requirement_scores'. [JD]: {jd_text} [Resume]: {json.dumps(resume_data)}"
        response = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(response_mime_type="application/json"))
        return json.loads(response.text)
    except: return None

# ---------------------------------------------------------
# PDF 生成與預覽邏輯 (全部使用獨立的暫存目錄避免名稱衝突)
# ---------------------------------------------------------
def generate_preview_pdf_bytes(data, template_name="main.tex", custom_tex_bytes=None, block_order=None):
    try:
        data_str = json.dumps(data, ensure_ascii=False)
        data_str = data_str.replace('**', '')
        clean_data = json.loads(data_str)

        with tempfile.TemporaryDirectory() as temp_dir:
            if custom_tex_bytes:
                template_name = "custom_main.tex"
                with open(os.path.join(temp_dir, template_name), "wb") as f:
                    f.write(custom_tex_bytes)
            else:
                shutil.copy(template_name, temp_dir)
            
            # --- Dynamic Block Injection ---
            tex_path = os.path.join(temp_dir, template_name)
            with open(tex_path, "r", encoding="utf-8") as f:
                template_content = f.read()

            if block_order and "BLOCKS_PLACEHOLDER" in template_content:
                blocks_latex = ""
                for block in block_order:
                    if block == "Summary":
                        blocks_latex += "\\directlua{printSummary()}\n"
                    elif block == "Experience":
                        if "elsa" in template_name:
                            blocks_latex += "\\section{PROFESSIONAL EXPERIENCE}\n  \\vspace{4pt}\n  \\directlua{printExperience()}\n"
                        else:
                            blocks_latex += "\\section{WORK EXPERIENCE}\n  \\directlua{printExperience()}\n"
                    elif block == "Education":
                        if "elsa" in template_name:
                            blocks_latex += "\\section{EDUCATION}\n  \\directlua{printEducation()}\n"
                        else:
                            blocks_latex += "\\section{EDUCATION}\n  \\directlua{printEducation()}\n"
                    elif block == "Projects & Patents":
                        blocks_latex += "\\directlua{printProjectsAndPatents()}\n"
                    elif block == "Skills":
                        if "elsa" in template_name:
                            blocks_latex += "\\section{CORE SKILLS}\n  \\directlua{printSkills()}\n"
                        else:
                            blocks_latex += "\\section{SKILLS}\n  \\directlua{printSkills()}\n"
                
                template_content = template_content.replace("BLOCKS_PLACEHOLDER", blocks_latex)
                with open(tex_path, "w", encoding="utf-8") as f:
                    f.write(template_content)
            # -------------------------------

            temp_json_path = os.path.join(temp_dir, "ml_resume.json")
            with open(temp_json_path, "w", encoding="utf-8") as f:
                json.dump(clean_data, f, ensure_ascii=False, indent=4)
                
            process = subprocess.run(
                ['lualatex', '-interaction=nonstopmode', template_name],
                cwd=temp_dir,
                capture_output=True
            )
            
            pdf_path = os.path.join(temp_dir, template_name.replace('.tex', '.pdf'))
            if process.returncode == 0 and os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    return f.read(), None
            else:
                return None, process.stdout.decode('utf-8', errors='ignore')
    except Exception as e:
        return None, str(e)

def generate_word_from_json(resume_data, block_order=None):
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)
        
    heading = resume_data.get("heading", {})
    p_name = doc.add_paragraph()
    p_name.alignment = 1
    run_name = p_name.add_run(heading.get("name", "Name"))
    run_name.bold = True
    run_name.font.size = Pt(20)
    
    contact_info = [heading.get("email", ""), heading.get("phone", ""), heading.get("linkedin", ""), heading.get("website", "")]
    doc.add_paragraph(" | ".join([c for c in contact_info if c])).alignment = 1
    
    if not block_order: block_order = ["Summary", "Experience", "Education", "Projects & Patents", "Skills"]
    for block in block_order:
        if block == "Summary" and resume_data.get("summary"):
            doc.add_heading("SUMMARY", level=1)
            doc.add_paragraph(resume_data.get("summary", ""))
        elif block == "Experience" and resume_data.get("experience"):
            doc.add_heading("WORK EXPERIENCE", level=1)
            for exp in resume_data.get("experience", []):
                p = doc.add_paragraph()
                p.add_run(exp.get("company", "")).bold = True
                p.add_run(f" - {exp.get('role', '')}").italic = True
                for d in exp.get("details", []): doc.add_paragraph(d.get("description", ""), style="List Bullet")
        elif block == "Education" and resume_data.get("education"):
            doc.add_heading("EDUCATION", level=1)
            for edu in resume_data.get("education", []):
                p = doc.add_paragraph()
                p.add_run(edu.get("school", "")).bold = True
                p.add_run(f" - {edu.get('degree', '')}")
        elif block == "Skills" and resume_data.get("skills"):
            doc.add_heading("SKILLS", level=1)
            for key in ["set1", "set2", "set3"]:
                s = resume_data["skills"].get(key, {})
                if s.get("title"): doc.add_paragraph(f"{s.get('title')}: {', '.join(s.get('items', []))}")
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

def generate_cover_letter_pdf_bytes(resume_data):
    try:
        cl_content = resume_data.get('cover_letter', '').replace('**', '')
        if not cl_content: return None, None, "No content."
        
        latex_template = r"""\documentclass[11pt]{article}\usepackage[margin=1in]{geometry}\usepackage{fontspec}\begin{document}""" + cl_content.replace("\n", "\n\n") + r"""\end{document}"""
        with tempfile.TemporaryDirectory() as temp_dir:
            with open(os.path.join(temp_dir, "cl.tex"), "w", encoding="utf-8") as f: f.write(latex_template)
            subprocess.run(['lualatex', '-interaction=nonstopmode', 'cl.tex'], cwd=temp_dir)
            with open(os.path.join(temp_dir, "cl.pdf"), "rb") as f: return f.read(), "cl.pdf", None
    except: return None, None, "Error"

def generate_cover_letter_word_bytes(resume_data):
    doc = Document()
    doc.add_paragraph(resume_data.get('cover_letter', '').replace('**', ''))
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

def render_pdf_js(pdf_bytes):
    base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
    pdf_js_html = f"""<html><head><script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script></head><body><div id="pdf-container"></div><script>pdfjsLib.GlobalWorkerOptions.workerSrc='https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';var binaryString=window.atob('{base64_pdf}');var bytes=new Uint8Array(binaryString.length);for(var i=0;i<binaryString.length;i++)bytes[i]=binaryString.charCodeAt(i);pdfjsLib.getDocument({{data:bytes}}).promise.then(function(pdf){{for(var i=1;i<=pdf.numPages;i++)pdf.getPage(i).then(function(page){{var canvas=document.createElement('canvas');document.getElementById('pdf-container').appendChild(canvas);page.render({{canvasContext:canvas.getContext('2d'),viewport:page.getViewport({{scale:1.5}})}});}});}});</script></body></html>"""
    st.html(pdf_js_html)

def get_glass_overlay_html(message="Processing...", animal_emoji="🐕"):
    return f"""<div style="position:fixed;top:0;left:0;width:100vw;height:100vh;background:rgba(0,0,0,0.7);backdrop-filter:blur(10px);z-index:9999;display:flex;justify-content:center;align-items:center;color:white;flex-direction:column;"><h1>{animal_emoji}</h1><h2>{message}</h2></div>"""

@st.dialog("Base Content", width="large")
def preview_base_profile(): st.json(st.session_state.resume_data)

# ---------------------------------------------------------
# Streamlit UI 介面
# ---------------------------------------------------------
st.set_page_config(page_title="AI Resume Builder", page_icon="🚀", layout="wide")

st.markdown("""<style>
.block-container { padding-top: 2rem; max-width: 1200px; }
div[data-testid="stVerticalBlock"] > div[style*="border"] { background-color: #ffffff05; border: 1px solid rgba(255, 255, 255, 0.1) !important; border-radius: 12px !important; padding: 1.5rem !important; }
.stButton > button { border-radius: 8px !important; height: 44px !important; transition: all 0.2s ease !important; }
.stButton > button[kind="primary"] { background: linear-gradient(135deg, #6366f1 0%, #a855f7 100%) !important; border: none !important; color: white !important; }
.stTabs [data-baseweb="tab-list"] { gap: 8px; }
.stTabs [data-baseweb="tab"] { background-color: #1e293b; border-radius: 8px 8px 0px 0px; padding: 0px 20px; }
.stTabs [aria-selected="true"] { border-bottom: 2px solid #6366f1 !important; }
[data-testid="stSidebar"] { background-color: #0f172a; }
</style>""", unsafe_allow_html=True)

db = init_firebase()

with st.sidebar:
    st.markdown("### 🚀 AI Resume Gen")
    st.markdown("---")
    if st.session_state.logged_in:
        st.success(f"Logged in: {st.session_state.user_email}")
        if st.button("Push to Cloud", use_container_width=True):
            save_user_profile(db, st.session_state.user_email, st.session_state.resume_data, st.session_state.custom_prompt, st.session_state.api_key)
        if st.button("Pull from Cloud", use_container_width=True):
            r, p, k = load_user_profile(db, st.session_state.user_email)
            if r: st.session_state.resume_data = r; st.session_state.custom_prompt = p; st.session_state.api_key = k; st.rerun()
        if st.button("Logout", use_container_width=True): st.session_state.logged_in = False; st.rerun()
    else:
        with st.form("login"):
            e = st.text_input("Email").strip(); p = st.text_input("Password", type="password")
            if st.form_submit_button("Login", type="primary", use_container_width=True):
                if authenticate_user(db, e, p): 
                    st.session_state.logged_in = True; st.session_state.user_email = e
                    r, pr, k = load_user_profile(db, e)
                    if r: st.session_state.resume_data = r; st.session_state.custom_prompt = pr; st.session_state.api_key = k
                    st.rerun()
    st.markdown("---")
    st.text_input("🔑 Gemini API Key", type="password", key="api_key")
    st.selectbox("🧠 Model", ["gemini-1.5-flash", "gemini-1.5-pro"], key="ai_model")
    st.color_picker("Brand Color", "#8a2be2", key="theme_color")
    st.selectbox("Animal", ["🦦 Otter", "🐕 Dog", "🦖 T-Rex"], key="animal_emoji_select")
    st.session_state.animal_emoji = st.session_state.animal_emoji_select.split(" ")[0]

st.title("🚀 Professional AI Resume")
s1, s2, s3, s4, s5 = len(st.session_state.resume_data.get("experience", [])) > 0, len(st.session_state.get("jd_input_v2", "")) > 50, st.session_state.optimized_resume_data is not None, st.session_state.resume_preview_bytes is not None, st.session_state.logged_in
steps = [{"l": "1. Source", "d": s1}, {"l": "2. Target", "d": s2}, {"l": "3. Analysis", "d": s3}, {"l": "4. Review", "d": s4}, {"l": "5. Tracker", "d": s5}]
cols = st.columns(5)
for i, s in enumerate(steps):
    with cols[i]:
        c = "#10b981" if s["d"] else "#6366f1"
        st.markdown(f"<div style='text-align:center;padding:10px;border-radius:10px;background:{c}15;border:1px solid {c}40;color:{c};font-weight:bold;'>{'✅' if s['d'] else '🔵'} {s['l']}</div>", unsafe_allow_html=True)

t1, t2, t3, t4, t5 = st.tabs([" 📁 Source ", " 🎯 Target ", " 📊 ATS Analysis ", " 📝 Editor & Export ", " 📈 Job Tracker "])

with t1:
    with st.container(border=True):
        st.subheader("📥 Quick Import")
        up = st.file_uploader("Upload PDF", type=["pdf"], key="up1")
        if st.button("✨ Extract", type="primary", use_container_width=True) and up:
            with st.status("Parsing..."):
                ok, msg, data = parse_pdf_resume_to_json(up.getvalue(), st.session_state.api_key)
                if ok: st.session_state.resume_data = data; st.rerun()
    st.markdown("#### 📝 Editor")
    edit = st_ace.st_ace(value=json.dumps(st.session_state.resume_data, indent=4, ensure_ascii=False), language="json", theme="dracula", height=500, key="base_ed")
    if st.button("💾 Save Base", use_container_width=True): st.session_state.resume_data = json.loads(edit); st.toast("Saved!")

with t2:
    with st.container(border=True):
        st.subheader("🎯 Job Details")
        jd = st.text_area("JD", height=300, key="jd_input_v2")
        st.text_area("Prompt", value=st.session_state.custom_prompt, key="custom_prompt_v2")
        c1, c2 = st.columns(2)
        with c1: ats = st.checkbox("ATS", value=True, key="ats_v2")
        with c2: visa = st.checkbox("Visa", value=True, key="visa_v2")
        st.markdown("---")
        co1, co2 = st.columns(2)
        with co1:
            if st.button("🚀 Start Optimization", type="primary", use_container_width=True):
                if jd:
                    l = st.empty(); l.markdown(get_glass_overlay_html("Optimizing...", st.session_state.animal_emoji), unsafe_allow_html=True)
                    ok, rep = ai_optimize_and_update(jd, st.session_state.custom_prompt_v2, ats, visa)
                    l.empty()
                    if ok: st.success("Done!")
                    else: st.error(f"Failed: {rep}")
        with co2:
            prompt = build_optimization_prompt(jd, st.session_state.custom_prompt_v2, ats, visa, st.session_state.resume_data)
            b64 = base64.b64encode(prompt.encode('utf-8')).decode('utf-8')
            st.html(f"""<button id="cp" onclick="navigator.clipboard.writeText(decodeURIComponent(escape(window.atob('{b64}')))).then(()=>{{this.innerText='✅ Copied';}})" style="width:100%;height:44px;border-radius:8px;background:#1e293b;color:white;border:1px solid #444;cursor:pointer;">📋 Copy Prompt</button>""")

with t3:
    st.header("📊 ATS Analysis")
    with st.expander("📥 Import External JSON"):
        ext = st.text_area("Paste JSON", height=200)
        if st.button("Apply"):
            try:
                d = json.loads(ext.strip('`json \n'))
                st.session_state.optimized_resume_data = d.get("optimized_resume", d)
                st.rerun()
            except: st.error("Invalid")
    if st.session_state.optimized_resume_data:
        if st.session_state.changelog: st.info(st.session_state.changelog)
        m = st.session_state.ats_metrics
        if m:
            c1, c2, c3 = st.columns(3)
            c1.metric("Match", f"{m['optimized_pct']}%")
            c2.metric("Hit", f"{m['optimized_count']}/{m['total']}")
            c3.metric("New", len(m['newly_added']))
            st.progress(m['optimized_pct']/100)

@st.dialog("🛠️ Tweak Data", width="large")
def edit_opt_dialog():
    edit = st_ace.st_ace(value=json.dumps(st.session_state.optimized_resume_data, indent=4, ensure_ascii=False), language="json", theme="dracula", height=500, auto_update=True)
    if st.button("💾 Save", use_container_width=True): 
        st.session_state.optimized_resume_data = json.loads(edit); st.rerun()

with t4:
    if st.session_state.optimized_resume_data:
        cl1, cl2 = st.columns([4, 6])
        with cl1:
            with st.container(border=True):
                if st.button("📝 Edit Data", use_container_width=True): edit_opt_dialog()
                tmpl = st.selectbox("Template", ["💻 Tech", "📈 Classic"], key="tmpl")
                order = st.multiselect("Order", ["Summary", "Experience", "Education", "Skills"], default=["Summary", "Experience", "Education", "Skills"])
                if st.button("🚀 Generate PDF", type="primary", use_container_width=True):
                    with st.spinner("Generating..."):
                        tex = "main.tex" if "Tech" in tmpl else "elsa_main.tex"
                        rb, _ = generate_preview_pdf_bytes(st.session_state.optimized_resume_data, tex, block_order=order)
                        if rb:
                            st.session_state.resume_preview_bytes = rb
                            c = st.session_state.optimized_resume_data.get('target_company','Company').replace(' ','_')
                            r = st.session_state.optimized_resume_data.get('target_role','Role').replace(' ','_')
                            st.session_state.resume_dl_data = {"bytes": rb, "name": f"{c}_{r}_Resume.pdf"}
                        cb, cn, _ = generate_cover_letter_pdf_bytes(st.session_state.optimized_resume_data)
                        if cb: st.session_state.cover_letter_preview_bytes = cb; st.session_state.cl_dl_data = {"bytes": cb, "name": f"{c}_{r}_CL.pdf"}
        with cl2:
            if st.session_state.resume_preview_bytes:
                type = st.radio("Display", ["Resume", "CL"], horizontal=True)
                data = st.session_state.resume_dl_data if type == "Resume" else st.session_state.cl_dl_data
                sync = st.checkbox("📈 Sync to Tracker", value=True) if st.session_state.logged_in else False
                if st.download_button(f"📥 Download {data['name']}", data["bytes"], data["name"], use_container_width=True):
                    if sync and type == "Resume": save_application(db, st.session_state.user_email, st.session_state.optimized_resume_data.get('target_company'), st.session_state.optimized_resume_data, st.session_state.get('jd_input_v2'))
                render_pdf_js(st.session_state.resume_preview_bytes if type == "Resume" else st.session_state.cover_letter_preview_bytes)
    else: st.warning("Optimize first.")

with t5:
    if st.session_state.logged_in: render_interview_progress(db, st.session_state.user_email); render_dashboard(db, st.session_state.user_email)
    else: st.warning("Login first.")
